import tkinter as tk
from tkinter import filedialog, ttk, messagebox
from pdf2docx import Converter
from pdf2image import convert_from_path
import pytesseract
from docx import Document
import os
from PIL import Image, ImageTk

class PDFToWordConverter:
    def __init__(self, root):
        self.root = root
        self.root.title("📄 PDF to Word Converter")
        self.root.geometry("600x450")
        self.root.minsize(600, 450)
        self.root.resizable(True, True)
        
        # สีธีมโมเดิร์น - Dark Mode
        self.colors = {
            'bg': '#1e1e1e',           # พื้นหลังหลัก
            'card': '#2d2d2d',         # การ์ด
            'primary': '#0078d4',      # สีหลัก
            'primary_hover': '#106ebe', # สีหลักเมื่อ hover
            'success': '#16a085',      # สีเขียว
            'warning': '#f39c12',      # สีเหลือง
            'error': '#e74c3c',        # สีแดง
            'text': '#ffffff',         # ข้อความหลัก
            'text_secondary': '#b0b0b0' # ข้อความรอง
        }
        
        self.root.configure(bg=self.colors['bg'])

        # ตั้งค่า theme
        style = ttk.Style()
        style.theme_use("clam")
        
        # ปรับแต่งสไตล์โมเดิร์น
        style.configure("Modern.TButton", 
                       padding=(20, 15), 
                       font=("Segoe UI", 12, "bold"), 
                       background=self.colors['primary'],
                       foreground="white",
                       borderwidth=0,
                       relief="flat")
        style.map("Modern.TButton", 
                 background=[("active", self.colors['primary_hover']),
                            ("pressed", "#005a9e")])
        
        style.configure("Card.TFrame", 
                       background=self.colors['card'],
                       relief="flat",
                       borderwidth=1)
        
        style.configure("Modern.TLabel", 
                       font=("Segoe UI", 11), 
                       background=self.colors['card'],
                       foreground=self.colors['text'],
                       anchor="center")
        
        style.configure("Title.TLabel", 
                       font=("Segoe UI", 24, "bold"), 
                       background=self.colors['bg'],
                       foreground=self.colors['text'],
                       anchor="center")
        
        style.configure("Modern.TProgressbar", 
                       thickness=8, 
                       troughcolor=self.colors['bg'],
                       background=self.colors['primary'],
                       borderwidth=0)

        # สร้าง frame หลัก
        self.main_frame = tk.Frame(self.root, bg=self.colors['bg'])
        self.main_frame.pack(fill="both", expand=True, padx=30, pady=30)
        
        # Title Section
        self.title_frame = tk.Frame(self.main_frame, bg=self.colors['bg'])
        self.title_frame.pack(fill="x", pady=(0, 30))
        
        self.label_title = tk.Label(self.title_frame, 
                                   text="📄 PDF to Word Converter", 
                                   font=("Segoe UI", 24, "bold"),
                                   bg=self.colors['bg'],
                                   fg=self.colors['text'])
        self.label_title.pack()
        
        self.subtitle = tk.Label(self.title_frame, 
                                text="แปลงไฟล์ PDF เป็น Word Document อย่างง่ายดาย", 
                                font=("Segoe UI", 10),
                                bg=self.colors['bg'],
                                fg=self.colors['text_secondary'])
        self.subtitle.pack(pady=(5, 0))
        
        # Main Content Card
        self.card_frame = tk.Frame(self.main_frame, 
                                  bg=self.colors['card'],
                                  relief="flat",
                                  bd=1)
        self.card_frame.pack(fill="both", expand=True, pady=(0, 20))
        
        # Add padding inside card
        self.content_frame = tk.Frame(self.card_frame, bg=self.colors['card'])
        self.content_frame.pack(fill="both", expand=True, padx=40, pady=40)

        # โหลดไอคอน (ใช้ emoji แทนไอคอน)
        self.icon_pdf = "📄"
        self.icon_convert = "🔄"
        self.icon_success = "✅"
        self.icon_error = "❌"
        
        # File Selection Section
        self.file_section = tk.Frame(self.content_frame, bg=self.colors['card'])
        self.file_section.pack(fill="x", pady=(0, 35))
        
        self.btn_select_pdf = tk.Button(self.file_section, 
                                       text=f"{self.icon_pdf} เลือกไฟล์ PDF", 
                                       font=("Segoe UI", 12, "bold"),
                                       bg=self.colors['primary'],
                                       fg="white",
                                       relief="flat",
                                       bd=0,
                                       padx=35,
                                       pady=18,
                                       cursor="hand2",
                                       command=self.select_pdf)
        self.btn_select_pdf.pack(pady=(0, 20))
        
        # เอฟเฟกต์ hover สำหรับปุ่ม
        self.btn_select_pdf.bind("<Enter>", lambda e: self.btn_select_pdf.config(bg=self.colors['primary_hover']))
        self.btn_select_pdf.bind("<Leave>", lambda e: self.btn_select_pdf.config(bg=self.colors['primary']))
        
        # File info display
        self.file_info_frame = tk.Frame(self.file_section, bg=self.colors['card'])
        self.file_info_frame.pack(fill="x")
        
        self.label_pdf = tk.Label(self.file_info_frame, 
                                 text="ยังไม่ได้เลือกไฟล์ PDF", 
                                 font=("Segoe UI", 10),
                                 bg=self.colors['card'],
                                 fg=self.colors['text_secondary'],
                                 wraplength=400)
        self.label_pdf.pack()
        
        # Convert Button
        self.btn_convert = tk.Button(self.content_frame, 
                                    text=f"{self.icon_convert} แปลงเป็น Word", 
                                    font=("Segoe UI", 13, "bold"),
                                    bg=self.colors['success'],
                                    fg="white",
                                    relief="flat",
                                    bd=0,
                                    padx=50,
                                    pady=20,
                                    cursor="hand2",
                                    state="disabled",
                                    command=self.convert_to_word)
        self.btn_convert.pack(pady=(0, 35))
        
        # เอฟเฟกต์ hover สำหรับปุ่มแปลง
        self.btn_convert.bind("<Enter>", self.on_convert_hover)
        self.btn_convert.bind("<Leave>", self.on_convert_leave)

        # Progress Section
        self.progress_frame = tk.Frame(self.content_frame, bg=self.colors['card'])
        self.progress_frame.pack(fill="x", pady=(0, 20))
        
        # สร้าง progress bar แบบกำหนดเอง
        self.progress_bg = tk.Frame(self.progress_frame, bg='#404040', height=8)
        self.progress_bg.pack(fill="x", padx=50, pady=10)
        
        self.progress_bar = tk.Frame(self.progress_bg, bg=self.colors['primary'], height=8)
        self.progress_bar.pack(side="left", fill="y")
        
        self.progress_frame.pack_forget()  # ซ่อนจนกว่าจะเริ่มแปลง
        
        # Animation variables
        self.progress_width = 0
        self.progress_direction = 1
        self.progress_animation = None
        
        # Status Section
        self.status_frame = tk.Frame(self.content_frame, bg=self.colors['card'])
        self.status_frame.pack(fill="x")
        
        self.label_status = tk.Label(self.status_frame, 
                                    text="", 
                                    font=("Segoe UI", 11),
                                    bg=self.colors['card'],
                                    fg=self.colors['text_secondary'])
        self.label_status.pack()

        self.pdf_path = None
    
    def animate_progress(self):
        """สร้างแอนิเมชัน progress bar"""
        if self.progress_animation:
            max_width = 300  # ความกว้างสูงสุด
            self.progress_width += self.progress_direction * 5
            
            if self.progress_width >= max_width:
                self.progress_direction = -1
            elif self.progress_width <= 0:
                self.progress_direction = 1
            
            self.progress_bar.config(width=self.progress_width)
            self.progress_animation = self.root.after(50, self.animate_progress)
    
    def start_progress(self):
        """เริ่มแอนิเมชัน progress"""
        self.progress_frame.pack(fill="x", pady=(0, 20))
        self.progress_animation = True
        self.animate_progress()
    
    def stop_progress(self):
        """หยุดแอนิเมชัน progress"""
        if self.progress_animation:
            self.root.after_cancel(self.progress_animation)
            self.progress_animation = None
        self.progress_frame.pack_forget()
        self.progress_width = 0
        self.progress_direction = 1
    
    def on_convert_hover(self, event):
        if self.btn_convert['state'] == 'normal':
            self.btn_convert.config(bg='#138d75')
    
    def on_convert_leave(self, event):
        if self.btn_convert['state'] == 'normal':
            self.btn_convert.config(bg=self.colors['success'])

    def select_pdf(self):
        self.pdf_path = filedialog.askopenfilename(filetypes=[("PDF files", "*.pdf")])
        if self.pdf_path:
            filename = os.path.basename(self.pdf_path)
            self.label_pdf.config(text=f"📄 {filename}", fg=self.colors['text'])
            self.btn_convert.config(state="normal", bg=self.colors['success'])
        else:
            self.label_pdf.config(text="ยังไม่ได้เลือกไฟล์ PDF", fg=self.colors['text_secondary'])
            self.btn_convert.config(state="disabled", bg='#555555')

    def convert_to_word(self):
        if not self.pdf_path:
            messagebox.showerror("ข้อผิดพลาด", "กรุณาเลือกไฟล์ PDF ก่อน")
            return

        docx_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if not docx_path:
            return

        # ตรวจสอบและเพิ่ม .docx ถ้าไม่มี
        if not docx_path.endswith(".docx"):
            docx_path += ".docx"

        self.label_status.config(text="🔄 กำลังแปลงไฟล์...", fg=self.colors['warning'])
        self.start_progress()
        self.root.update()

        try:
            # ลองแปลงด้วย pdf2docx ก่อน
            cv = Converter(self.pdf_path)
            cv.convert(docx_path, start=0, end=None)
            cv.close()

            # ตรวจสอบว่า PDF อาจเป็นภาพ (ถ้าไฟล์ Word ว่างเปล่า)
            doc = Document(docx_path)
            if not doc.paragraphs or not any(p.text.strip() for p in doc.paragraphs):
                self.label_status.config(text="🔍 กำลังใช้ OCR สำหรับไฟล์ภาพ...", fg=self.colors['warning'])
                self.root.update()
                self.ocr_conversion(self.pdf_path, docx_path)

            self.stop_progress()
            self.label_status.config(text=f"{self.icon_success} แปลงไฟล์เสร็จสิ้น!", fg=self.colors['success'])
            messagebox.showinfo("สำเร็จ", f"แปลงไฟล์เป็น {os.path.basename(docx_path)} เรียบร้อยแล้ว")

        except Exception as e:
            self.stop_progress()
            self.label_status.config(text=f"{self.icon_error} เกิดข้อผิดพลาดในการแปลง", fg=self.colors['error'])
            messagebox.showerror("ข้อผิดพลาด", f"เกิดข้อผิดพลาด: {str(e)}")

    def ocr_conversion(self, pdf_path, docx_path):
        # แปลง PDF เป็นภาพ
        images = convert_from_path(pdf_path)
        doc = Document()

        # ดึงข้อความจากแต่ละภาพด้วย Tesseract
        for image in images:
            text = pytesseract.image_to_string(image, lang="tha+eng")  # รองรับภาษาไทยและอังกฤษ
            doc.add_paragraph(text)

        # บันทึกไฟล์ Word
        doc.save(docx_path)

if __name__ == "__main__":
    root = tk.Tk()
    app = PDFToWordConverter(root)
    root.mainloop()
